#!/usr/bin/env python3
"""
Conversión del Manual de Usuario de Anclora Nexus a formato DOCX
Implementa la skill manual-format-exporter.md
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
from pathlib import Path
import re
import os

# Colores de marca Anclora
NAVY = RGBColor(25, 35, 80)  # #192350
GOLD = RGBColor(212, 175, 55)  # #D4AF37
WHITE_SOFT = RGBColor(245, 245, 240)  # #F5F5F0

class DocxManualExporter:
    def __init__(self, markdown_path: str, output_path: str):
        self.markdown_path = markdown_path
        self.output_path = output_path
        self.doc = Document()

    def setup_styles(self):
        """Configura los estilos del documento siguiendo la marca Anclora"""
        styles = self.doc.styles

        # Estilo para Título 1 (Playfair Display simulado con serif)
        h1 = styles['Heading 1']
        h1.font.name = 'Georgia'  # Aproximación a Playfair Display
        h1.font.size = Pt(24)
        h1.font.color.rgb = NAVY
        h1.font.bold = True

        # Estilo para Título 2
        h2 = styles['Heading 2']
        h2.font.name = 'Georgia'
        h2.font.size = Pt(18)
        h2.font.color.rgb = GOLD
        h2.font.bold = True

        # Estilo para Título 3
        h3 = styles['Heading 3']
        h3.font.name = 'Arial'
        h3.font.size = Pt(14)
        h3.font.color.rgb = NAVY
        h3.font.bold = True

        # Estilo para párrafos normales (Inter simulado con Arial)
        normal = styles['Normal']
        normal.font.name = 'Arial'
        normal.font.size = Pt(11)
        normal.font.color.rgb = RGBColor(0, 0, 0)

    def add_cover_page(self, logo_path: str = None):
        """Añade portada con branding Anclora"""
        # Logo (si existe)
        if logo_path and os.path.exists(logo_path):
            self.doc.add_paragraph()  # Spacing
            logo_para = self.doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = logo_para.add_run()
            run.add_picture(logo_path, width=Inches(3))
            self.doc.add_paragraph()  # Spacing
        else:
            # Si no hay logo, más espaciado superior
            for _ in range(3):
                self.doc.add_paragraph()

        # Título principal
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('ANCLORA NEXUS')
        run.font.name = 'Georgia'
        run.font.size = Pt(36)
        run.font.color.rgb = NAVY
        run.font.bold = True

        # Línea decorativa
        line = self.doc.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = line.add_run('━━━━━━━━━━━━━━━━━')
        run.font.color.rgb = GOLD
        run.font.size = Pt(14)

        # Subtítulo
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('MANUAL DE USUARIO')
        run.font.name = 'Arial'
        run.font.size = Pt(24)
        run.font.color.rgb = GOLD
        run.font.bold = True

        # Descripción
        desc = self.doc.add_paragraph()
        desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = desc.add_run('\nCapa de Inteligencia Inmobiliaria\nPara Agentes Independientes')
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.color.rgb = NAVY

        # Versión y fecha
        self.doc.add_paragraph('\n' * 8)
        version_para = self.doc.add_paragraph()
        version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = version_para.add_run(f'Versión 1.2.3\n{datetime.now().strftime("%d/%m/%Y")}')
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)

        # Salto de página
        self.doc.add_page_break()

    def add_toc(self):
        """Añade índice de contenidos"""
        toc_heading = self.doc.add_heading('Tabla de Contenidos', level=1)
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        note = self.doc.add_paragraph()
        run = note.add_run('Nota: El índice se actualizará automáticamente al abrir el documento en Microsoft Word.')
        run.font.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)

        self.doc.add_page_break()

    def add_screenshot(self, image_path: str, caption: str = None):
        """Add screenshot to DOCX"""
        if Path(image_path).exists():
            try:
                # Add image (max width 6 inches to fit page)
                self.doc.add_picture(image_path, width=Inches(6))

                # Add caption if provided
                if caption:
                    caption_para = self.doc.add_paragraph()
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = caption_para.add_run(f"Figura: {caption}")
                    run.font.size = Pt(9)
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(100, 100, 100)

                self.doc.add_paragraph()  # Spacing
                return True
            except Exception as e:
                print(f"   ⚠️  Error añadiendo screenshot {image_path}: {e}")
                return False
        return False

    def parse_markdown_to_docx(self, markdown_content: str):
        """Convierte markdown a formato DOCX"""
        lines = markdown_content.split('\n')
        in_code_block = False
        in_list = False
        list_content = []

        # Screenshot mapping (relative to base_dir)
        base_dir = "/home/dev/proyectos/anclora-nexus"
        screenshot_map = {
            "### 3.1 Dashboard": f"{base_dir}/public/docs/manual-usuario/assets/screenshots/01-dashboard.png",
            "### 3.2 Leads": f"{base_dir}/public/docs/manual-usuario/assets/screenshots/02-leads.png",
            "### 3.3 Properties": f"{base_dir}/public/docs/manual-usuario/assets/screenshots/03-properties.png",
            "### 3.4 Tasks": f"{base_dir}/public/docs/manual-usuario/assets/screenshots/04-tasks.png",
            "### 3.5 Team": f"{base_dir}/public/docs/manual-usuario/assets/screenshots/05-team.png",
            "### 4.2 Prospection operativa": f"{base_dir}/public/docs/manual-usuario/assets/screenshots/06-prospection-unified.png",
            "### 4.3 Seller Pipeline": f"{base_dir}/public/docs/manual-usuario/assets/screenshots/07-sellers.png",
            "### 4.4 Opportunity Ranking": f"{base_dir}/public/docs/manual-usuario/assets/screenshots/08-opportunity-ranking.png",
            "### 4.5 Intelligence": f"{base_dir}/public/docs/manual-usuario/assets/screenshots/09-intelligence.png",
            "### 5.1 Ingestion": f"{base_dir}/public/docs/manual-usuario/assets/screenshots/11-ingestion.png",
            "### 5.2 Data Quality": f"{base_dir}/public/docs/manual-usuario/assets/screenshots/12-data-quality.png",
            "### 5.3 Feed Orchestrator": f"{base_dir}/public/docs/manual-usuario/assets/screenshots/13-feed-orchestrator.png",
            "### 5.4 Automation & Alerting": f"{base_dir}/public/docs/manual-usuario/assets/screenshots/14-automation-alerting.png",
            "### 5.5 Command Center": f"{base_dir}/public/docs/manual-usuario/assets/screenshots/15-command-center.png",
            "### 5.6 Deal Margin Simulator": f"{base_dir}/public/docs/manual-usuario/assets/screenshots/16-deal-margin-simulator.png",
            "### 5.7 Source Observatory": f"{base_dir}/public/docs/manual-usuario/assets/screenshots/17-source-observatory.png",
        }

        for line in lines:
            # Ignorar frontmatter YAML
            if line.strip() == '---':
                continue
            if line.startswith('title:') or line.startswith('version:') or line.startswith('date:') or line.startswith('language:') or line.startswith('status:'):
                continue

            # Código en bloque
            if line.strip().startswith('```'):
                in_code_block = not in_code_block
                continue

            if in_code_block:
                p = self.doc.add_paragraph(line)
                p.style = 'Normal'
                p.paragraph_format.left_indent = Inches(0.5)
                run = p.runs[0]
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(60, 60, 60)
                continue

            # Títulos
            if line.startswith('# ') and not line.startswith('## '):
                self.doc.add_heading(line[2:], level=1)
                in_list = False
            elif line.startswith('## '):
                heading_text = line[3:]
                self.doc.add_heading(heading_text, level=2)
                in_list = False

                # Check if we should insert a screenshot after this heading
                if line in screenshot_map:
                    screenshot_path = screenshot_map[line]
                    if Path(screenshot_path).exists():
                        self.add_screenshot(screenshot_path, heading_text)

            elif line.startswith('### '):
                heading_text = line[4:]
                self.doc.add_heading(heading_text, level=3)
                in_list = False

                # Check if we should insert a screenshot after this heading
                if line in screenshot_map:
                    screenshot_path = screenshot_map[line]
                    if Path(screenshot_path).exists():
                        self.add_screenshot(screenshot_path, heading_text)

            # Listas
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                content = line.strip()[2:]
                p = self.doc.add_paragraph(content, style='List Bullet')
                in_list = True

            # Listas numeradas
            elif re.match(r'^\d+\.\s', line.strip()):
                content = re.sub(r'^\d+\.\s', '', line.strip())
                p = self.doc.add_paragraph(content, style='List Number')
                in_list = True

            # Tablas (formato simple)
            elif '|' in line and line.strip():
                if not line.strip().startswith('|---'):
                    cells = [cell.strip() for cell in line.split('|')[1:-1]]
                    if cells:
                        # Detectar si es encabezado
                        if re.match(r'^[A-Z\s]+$', ''.join(cells)):
                            p = self.doc.add_paragraph()
                            for i, cell in enumerate(cells):
                                run = p.add_run(cell)
                                run.font.bold = True
                                run.font.color.rgb = NAVY
                                if i < len(cells) - 1:
                                    p.add_run(' | ')
                        else:
                            p = self.doc.add_paragraph(' | '.join(cells))
                            p.paragraph_format.left_indent = Inches(0.25)
                in_list = False

            # Texto enfatizado
            elif line.strip():
                in_list = False
                p = self.doc.add_paragraph()

                # Procesar negrita **texto**
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.font.bold = True
                    else:
                        # Procesar cursiva *texto*
                        italic_parts = re.split(r'(\*.*?\*)', part)
                        for ipart in italic_parts:
                            if ipart.startswith('*') and ipart.endswith('*') and not ipart.startswith('**'):
                                run = p.add_run(ipart[1:-1])
                                run.font.italic = True
                            else:
                                p.add_run(ipart)

            # Líneas vacías
            else:
                if not in_list:
                    self.doc.add_paragraph()

    def add_footer(self):
        """Añade pie de página con marca y numeración"""
        section = self.doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = footer_para.add_run('Anclora Nexus — Intelligence Layer')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)

    def export(self, logo_path: str = None):
        """Ejecuta la exportación completa"""
        print(f"🔄 Leyendo markdown: {self.markdown_path}")

        with open(self.markdown_path, 'r', encoding='utf-8') as f:
            markdown_content = f.read()

        print("🎨 Configurando estilos de marca Anclora...")
        self.setup_styles()

        print("📄 Generando portada...")
        if logo_path and os.path.exists(logo_path):
            print(f"   🎨 Añadiendo logo: {logo_path}")
        self.add_cover_page(logo_path)

        print("📑 Añadiendo tabla de contenidos...")
        self.add_toc()

        print("✍️  Convirtiendo contenido markdown a DOCX...")
        self.parse_markdown_to_docx(markdown_content)

        print("🔖 Añadiendo pie de página...")
        self.add_footer()

        print(f"💾 Guardando documento: {self.output_path}")
        self.doc.save(self.output_path)

        # Verificar tamaño
        size = os.path.getsize(self.output_path)
        print(f"✅ Documento generado exitosamente!")
        print(f"   Ruta: {self.output_path}")
        print(f"   Tamaño: {size:,} bytes ({size/1024:.1f} KB)")


if __name__ == "__main__":
    import sys

    # Parse language argument
    lang = "es"  # default
    if len(sys.argv) > 1:
        if sys.argv[1] == "--lang" and len(sys.argv) > 2:
            lang = sys.argv[2].lower()

    # Rutas
    base_dir = "/home/dev/proyectos/anclora-nexus"

    if lang == "en":
        markdown_path = f"{base_dir}/public/docs/manual-usuario/MANUAL_USUARIO_ANCLORA_NEXUS_EN.md"
        docx_path = f"{base_dir}/public/docs/manual-usuario/MANUAL_USUARIO_ANCLORA_NEXUS_EN.docx"
        print_lang = "EN (ENGLISH)"
    else:
        markdown_path = f"{base_dir}/public/docs/manual-usuario/MANUAL_USUARIO_ANCLORA_NEXUS.md"
        docx_path = f"{base_dir}/public/docs/manual-usuario/MANUAL_USUARIO_ANCLORA_NEXUS.docx"
        print_lang = "ES (ESPAÑOL)"

    logo_path = f"{base_dir}/public/brand/anclora-nexus.png"

    # Ejecutar conversión
    print("=" * 60)
    print(f"🚀 CONVERSIÓN DE MANUAL DE USUARIO A DOCX ({print_lang})")
    print("=" * 60)

    exporter = DocxManualExporter(markdown_path, docx_path)
    exporter.export(logo_path=logo_path)

    print("\n" + "=" * 60)
    print("✨ Conversión completada!")
    print("=" * 60)
    print(f"\n💡 Para generar otra versión: python3 {sys.argv[0]} --lang [es|en]")
