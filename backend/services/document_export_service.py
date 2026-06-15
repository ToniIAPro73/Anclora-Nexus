"""Professional document export: markdown → PDF (weasyprint) and DOCX (python-docx)."""

from __future__ import annotations

import base64
import io
import os
import re
from typing import Any

import markdown as _md

LOGO_PATH = os.path.join(os.path.dirname(__file__), "../assets/logo.png")

# ── Brand CSS ─────────────────────────────────────────────────────────────────

_CSS = """
@font-face {
    font-family: 'Inter';
    src: local('Inter'), local('Arial');
}

* { box-sizing: border-box; }

body {
    font-family: Georgia, 'Times New Roman', Times, serif;
    font-size: 10.5pt;
    color: #1a1a2e;
    line-height: 1.65;
    margin: 0;
    background: #fff;
}

.doc-wrapper {
    padding: 0 10mm;
}

.logo-block {
    text-align: center;
    margin-bottom: 18pt;
    padding-bottom: 14pt;
    border-bottom: 2px solid #c9a84c;
}

.logo-block img {
    height: 48pt;
    object-fit: contain;
}

h1 {
    font-family: Georgia, serif;
    font-size: 18pt;
    color: #0f1629;
    text-align: center;
    margin: 18pt 0 8pt;
    font-weight: 700;
    letter-spacing: 0.02em;
}

h2 {
    font-family: Georgia, serif;
    font-size: 11.5pt;
    color: #0f1629;
    margin-top: 22pt;
    margin-bottom: 8pt;
    padding-bottom: 5pt;
    border-bottom: 1px solid #ddd5bb;
    font-weight: 700;
    letter-spacing: 0.01em;
}

h3 {
    font-family: Georgia, serif;
    font-size: 10.5pt;
    color: #2a3765;
    margin-top: 14pt;
    margin-bottom: 5pt;
    font-weight: 700;
}

p {
    margin: 5pt 0 8pt;
    text-align: justify;
    hyphens: auto;
}

hr {
    border: none;
    border-top: 1px solid #c9a84c;
    margin: 16pt 0;
}

/* Tables */
table {
    width: 100%;
    border-collapse: collapse;
    margin: 10pt 0 14pt;
    font-size: 9.5pt;
}

thead tr {
    background-color: #0f1629;
    color: #fff;
}

th {
    padding: 7pt 10pt;
    text-align: left;
    font-weight: 600;
    font-family: Arial, sans-serif;
    font-size: 9pt;
    letter-spacing: 0.03em;
    text-transform: uppercase;
}

td {
    padding: 6pt 10pt;
    border-bottom: 1px solid #e8e0cc;
    vertical-align: top;
    color: #222;
}

tr:nth-child(even) td {
    background-color: #faf8f3;
}

/* Lists */
ul, ol {
    margin: 6pt 0 8pt 18pt;
    padding: 0;
}

li {
    margin: 3pt 0;
    padding-left: 4pt;
}

strong {
    color: #0f1629;
    font-weight: 700;
}

em { font-style: italic; }

/* Signature block */
.signature-row {
    display: flex;
    justify-content: space-between;
    margin-top: 32pt;
    gap: 40pt;
}

/* Page layout */
@page {
    size: A4;
    margin: 20mm 18mm 22mm 18mm;

    @top-center {
        content: element(running-header);
    }

    @bottom-center {
        content: "";
        border-top: 0.5px solid #c9a84c;
        width: 90%;
        margin: 0 auto;
    }

    @bottom-right {
        content: counter(page) " / " counter(pages);
        font-family: Arial, sans-serif;
        font-size: 8pt;
        color: #888;
        margin-bottom: 4mm;
    }

    @bottom-left {
        content: "Anclora Private Estates S.L. — Documento confidencial";
        font-family: Arial, sans-serif;
        font-size: 8pt;
        color: #888;
        margin-bottom: 4mm;
    }
}
"""


# ── Logo helper ───────────────────────────────────────────────────────────────

def _logo_data_uri() -> str | None:
    """Return the logo as a data URI for embedding in HTML."""
    try:
        with open(LOGO_PATH, "rb") as f:
            b64 = base64.b64encode(f.read()).decode()
        return f"data:image/png;base64,{b64}"
    except FileNotFoundError:
        return None


# ── Markdown → HTML ───────────────────────────────────────────────────────────

def _md_to_html_body(md_text: str, logo_uri: str | None) -> str:
    """Convert markdown text to a styled HTML document body."""
    # Replace logo marker with embedded logo block
    if logo_uri:
        logo_block = f'<div class="logo-block"><img src="{logo_uri}" alt="Anclora Private Estates"/></div>'
    else:
        logo_block = ""
    md_text = re.sub(r"!\[.*?\]\(ANCLORA_LOGO_PLACEHOLDER\)", logo_block, md_text)

    # Strip stray raw markdown image refs that weren't replaced
    md_text = re.sub(r"!\[.*?\]\([^)]*PLACEHOLDER[^)]*\)", "", md_text)

    body_html = _md.markdown(
        md_text,
        extensions=["tables", "nl2br", "attr_list", "sane_lists"],
    )
    return body_html


def _build_html(md_text: str) -> str:
    logo_uri = _logo_data_uri()
    body = _md_to_html_body(md_text, logo_uri)
    return f"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="utf-8"/>
<style>{_CSS}</style>
</head>
<body>
<div class="doc-wrapper">
{body}
</div>
</body>
</html>"""


# ── PDF export ────────────────────────────────────────────────────────────────

def export_pdf(md_text: str) -> bytes:
    """Render markdown → professional PDF via weasyprint."""
    try:
        import weasyprint  # noqa: PLC0415
        html_src = _build_html(md_text)
        return weasyprint.HTML(string=html_src).write_pdf()
    except Exception as exc:
        raise RuntimeError(f"PDF generation failed: {exc}") from exc


# ── DOCX export ───────────────────────────────────────────────────────────────

def export_docx(md_text: str) -> bytes:
    """Render markdown → professional DOCX via python-docx."""
    try:
        from docx import Document  # noqa: PLC0415
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: PLC0415
        from docx.oxml import OxmlElement  # noqa: PLC0415
        from docx.oxml.ns import qn  # noqa: PLC0415
        from docx.shared import Inches, Pt, RGBColor  # noqa: PLC0415

        doc = Document()

        # ── Page margins (A4, narrow) ──────────────────────────────────────
        section = doc.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(1.18)
        section.top_margin = Inches(0.98)
        section.bottom_margin = Inches(0.98)

        # ── Logo in header ─────────────────────────────────────────────────
        if os.path.exists(LOGO_PATH):
            header = section.header
            hpara = header.paragraphs[0]
            hpara.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = hpara.add_run()
            run.add_picture(LOGO_PATH, height=Pt(28))
            # Gold separator line under header
            _add_paragraph_border(hpara, bottom=True)

        # ── Styles ────────────────────────────────────────────────────────
        _configure_styles(doc)

        # ── Parse and build ───────────────────────────────────────────────
        # Remove logo placeholder lines — already in header
        md_text = re.sub(r"!\[.*?\]\(ANCLORA_LOGO_PLACEHOLDER\)\n?", "", md_text)

        _parse_markdown_to_doc(doc, md_text)

        # ── Serialize ─────────────────────────────────────────────────────
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception as exc:
        raise RuntimeError(f"DOCX generation failed: {exc}") from exc


def _configure_styles(doc: Any) -> None:
    """Apply brand typography to Word built-in styles."""
    from docx.shared import Pt, RGBColor  # noqa: PLC0415

    navy = RGBColor(0x0F, 0x16, 0x29)
    gold = RGBColor(0xC9, 0xA8, 0x4C)

    styles_cfg = {
        "Heading 1": {"size": 20, "bold": True, "color": navy, "space_before": 12, "space_after": 6},
        "Heading 2": {"size": 13, "bold": True, "color": navy, "space_before": 18, "space_after": 4},
        "Heading 3": {"size": 11, "bold": True, "color": RGBColor(0x2A, 0x37, 0x65), "space_before": 12, "space_after": 3},
        "Normal":    {"size": 10, "color": RGBColor(0x1A, 0x1A, 0x2E), "space_after": 4},
    }

    for style_name, cfg in styles_cfg.items():
        try:
            style = doc.styles[style_name]
            font = style.font
            font.size = Pt(cfg["size"])
            font.bold = cfg.get("bold", False)
            font.color.rgb = cfg["color"]
            pf = style.paragraph_format
            if "space_before" in cfg:
                pf.space_before = Pt(cfg["space_before"])
            if "space_after" in cfg:
                pf.space_after = Pt(cfg["space_after"])
        except Exception:
            pass


def _add_paragraph_border(para: Any, bottom: bool = False) -> None:
    """Add a decorative border to a paragraph via XML."""
    from docx.oxml import OxmlElement  # noqa: PLC0415
    from docx.oxml.ns import qn  # noqa: PLC0415

    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    if bottom:
        bdr = OxmlElement("w:bottom")
        bdr.set(qn("w:val"), "single")
        bdr.set(qn("w:sz"), "6")
        bdr.set(qn("w:space"), "1")
        bdr.set(qn("w:color"), "C9A84C")
        pBdr.append(bdr)
    pPr.append(pBdr)


def _parse_markdown_to_doc(doc: Any, md_text: str) -> None:
    """Parse markdown line by line and add content to the Word document."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: PLC0415
    from docx.shared import Pt, RGBColor  # noqa: PLC0415

    lines = md_text.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]

        # ── Horizontal rule ───────────────────────────────────────────────
        if re.match(r"^---+$", line.strip()):
            para = doc.add_paragraph()
            _add_paragraph_border(para, bottom=True)
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # ── ATX Headings ──────────────────────────────────────────────────
        m = re.match(r"^(#{1,3})\s+(.+)$", line)
        if m:
            level = len(m.group(1))
            text = _strip_inline(m.group(2))
            style = f"Heading {level}"
            para = doc.add_paragraph(text, style=style)
            if level == 2:
                _add_paragraph_border(para, bottom=True)
            i += 1
            continue

        # ── Table ─────────────────────────────────────────────────────────
        if "|" in line and i + 1 < len(lines) and re.match(r"^\|?\s*[-:]+", lines[i + 1]):
            # Collect all table rows
            table_lines: list[str] = []
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i])
                i += 1
            _add_table(doc, table_lines)
            continue

        # ── Unordered list ────────────────────────────────────────────────
        if re.match(r"^[-*+]\s+", line):
            text = _strip_inline(re.sub(r"^[-*+]\s+", "", line))
            para = doc.add_paragraph(text, style="List Bullet")
            i += 1
            continue

        # ── Ordered list ──────────────────────────────────────────────────
        if re.match(r"^\d+\.\s+", line):
            text = _strip_inline(re.sub(r"^\d+\.\s+", "", line))
            doc.add_paragraph(text, style="List Number")
            i += 1
            continue

        # ── Empty line ────────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── Normal paragraph ─────────────────────────────────────────────
        para = doc.add_paragraph()
        _add_inline_runs(para, line)
        i += 1


def _add_table(doc: Any, table_lines: list[str]) -> None:
    """Build a Word table from markdown table lines."""
    from docx.shared import Pt, RGBColor  # noqa: PLC0415

    def split_row(line: str) -> list[str]:
        cells = line.strip().strip("|").split("|")
        return [c.strip() for c in cells]

    rows = [split_row(l) for l in table_lines if not re.match(r"^\|?\s*[-:]+", l)]
    if not rows:
        return

    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"

    navy = RGBColor(0x0F, 0x16, 0x29)

    for r_idx, row_cells in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_cells):
            if c_idx >= col_count:
                break
            cell = row.cells[c_idx]
            para = cell.paragraphs[0]
            run = para.add_run(_strip_inline(cell_text))
            run.font.size = Pt(9.5)
            if r_idx == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                # Set header cell background
                try:
                    from docx.oxml import OxmlElement  # noqa: PLC0415
                    from docx.oxml.ns import qn  # noqa: PLC0415
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:fill"), "0F1629")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:val"), "clear")
                    tcPr.append(shd)
                except Exception:
                    pass


def _strip_inline(text: str) -> str:
    """Remove markdown inline markers from text."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\([^)]*\)", r"\1", text)
    return text.strip()


def _add_inline_runs(para: Any, text: str) -> None:
    """Add inline-styled runs (bold, italic) to a paragraph."""
    from docx.shared import RGBColor  # noqa: PLC0415

    navy = RGBColor(0x0F, 0x16, 0x29)
    # Split by bold markers
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
            run.font.color.rgb = navy
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            if part:
                para.add_run(part)
